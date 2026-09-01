"""文档解析器工厂（Task 4）：按扩展名分发，文件字节 → list[langchain Document]。

- pdf：优先 LangChain PyPDFLoader（底层 pypdf），不可用时直接用 pypdf 逐页解析；
- docx：python-docx 按标题分节聚合，同一章节的段落与表格行（"表头: 值"行文本）
  合并为一个 Document，标题文本进内容，title_path 记入 metadata；
- xls/xlsx：openpyxl 逐工作表，每行转 "表头1: 值1; 表头2: 值2..."，sheet 名入 metadata；
- txt/md：直接读文本（md 保留原文，切分阶段用 MarkdownHeaderTextSplitter）；
- .doc（老格式）：不支持，置 FAILED 并提示转存 .docx（由任务层捕获 UnsupportedFormatError）。
"""

import io
import os
import re
import tempfile

from langchain_core.documents import Document

from app.core.log import get_logger
from app.document.cleaner import clean_documents, normalize_text

logger = get_logger(__name__)

# 上传/解析支持的扩展名（.doc 会被显式拒绝并给出明确提示）
SUPPORTED_EXTENSIONS = {"pdf", "doc", "docx", "xls", "xlsx", "txt", "md"}

# docx 标题样式名（兼容中英文 Word："Heading 1" / "标题 1"）
_HEADING_STYLE_RE = re.compile(r"^(?:Heading|标题)\s*([1-9])$", re.IGNORECASE)


class UnsupportedFormatError(Exception):
    """文件格式不支持解析（任务层据此置 FAILED）。"""


def file_extension(filename: str) -> str:
    """取小写扩展名（无扩展名返回空串）。"""
    name = (filename or "").rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


def _decode_text(content: bytes) -> str:
    """文本解码：utf-8 优先，其次 gb18030，兜底替换非法字节。"""
    for encoding in ("utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


# ----- PDF -----

# 标题判定：行字号 ≥ 正文字号 × 该倍数、行长受限、不以句末标点结尾
_HEADING_SIZE_RATIO = 1.15
_HEADING_MAX_LEN = 50
_TRAILING_PUNCT = ("。", "，", "；", "、", ",")
# 标题字号聚类容差（pt）：相近字号视为同一层级
_LEVEL_TOLERANCE = 0.5
# 页面行类型：text=普通文本行（参与字号统计/标题判定）；table=表格 markdown 行；
# ocr=OCR 识别行（均只作正文，不参与标题推断）
_KIND_TEXT, _KIND_TABLE, _KIND_OCR = "text", "table", "ocr"

# OCR 引擎进程级单例（None=未初始化，False=不可用）
_OCR_ENGINE = None


def _get_ocr_engine():
    """RapidOCR 引擎单例；未安装返回 None（OCR 静默跳过）。模型随包分发，离线可用。"""
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        try:
            from rapidocr_onnxruntime import RapidOCR
        except Exception as exc:
            logger.warning("rapidocr unavailable, pdf ocr disabled: %s", exc)
            _OCR_ENGINE = False
            return None
        _OCR_ENGINE = RapidOCR()
    return _OCR_ENGINE or None


def _ocr_lines(engine, img) -> list[str]:
    """对图像（BGR ndarray）做 OCR，返回按版面顺序（先上后下、先左后右）的行文本。"""
    result, _ = engine(img)
    if not result:
        return []
    items = sorted(result, key=lambda r: (r[0][0][1], r[0][0][0]))
    return [str(r[1]).strip() for r in items if str(r[1]).strip()]


def _ocr_page(page, engine) -> list[str]:
    """整页渲染后 OCR（扫描页）。"""
    import cv2
    import numpy as np

    pix = page.get_pixmap(dpi=200)
    img = cv2.imdecode(np.frombuffer(pix.tobytes("png"), np.uint8), cv2.IMREAD_COLOR)
    return _ocr_lines(engine, img) if img is not None else []


def _ocr_embedded_images(doc, page, engine) -> list[tuple[float, list[str]]]:
    """提取页面内嵌图片逐张 OCR，返回 [(图片 y0, 行文本列表)]。"""
    import cv2
    import numpy as np

    out: list[tuple[float, list[str]]] = []
    for img_info in page.get_images(full=True):
        xref = img_info[0]
        try:
            raw = doc.extract_image(xref)["image"]
            img = cv2.imdecode(np.frombuffer(raw, np.uint8), cv2.IMREAD_COLOR)
            if img is None:
                continue
            lines = _ocr_lines(engine, img)
            if lines:
                rects = page.get_image_rects(xref)
                out.append((rects[0].y0 if rects else 0.0, lines))
        except Exception as exc:
            logger.debug("ocr embedded image xref=%s failed: %s", xref, exc)
    return out


def _parse_pdf(content: bytes) -> list[Document]:
    """PDF：优先 PyMuPDF 结构化解析（字号推断标题 + 分节聚合），不可用回退纯文本。"""
    try:
        import pymupdf
    except Exception as exc:  # pragma: no cover - 依赖缺失时回退
        logger.warning("pymupdf unavailable (%s), fallback to plaintext pdf parse", exc)
        return _parse_pdf_plaintext(content)
    try:
        return _parse_pdf_structured(content, pymupdf)
    except Exception as exc:
        logger.warning("structured pdf parse failed (%s), fallback to plaintext", exc)
        return _parse_pdf_plaintext(content)


def _parse_pdf_structured(content: bytes, pymupdf) -> list[Document]:
    """PyMuPDF 解析：表格转 markdown + 图片 OCR → 页眉页脚剔除 → 字号推断标题 → 分节聚合。"""
    from collections import Counter

    from app.core.config import settings
    from app.document.cleaner import (  # noqa: PLC0415
        _EDGE_LINES,
        _is_noise_line,
        find_repeated_edge_lines,
    )
    from app.document.sections import SectionAggregator

    def _line_in_table(line_bbox: tuple, table_bboxes: list[tuple]) -> bool:
        """行中心点落在任一表格 bbox 内视为表格内部行（由 markdown 块替代）。"""
        cx = (line_bbox[0] + line_bbox[2]) / 2
        cy = (line_bbox[1] + line_bbox[3]) / 2
        return any(x0 <= cx <= x1 and y0 <= cy <= y1 for x0, y0, x1, y1 in table_bboxes)

    ocr = _get_ocr_engine() if settings.PDF_OCR_ENABLED else None

    doc = pymupdf.open(stream=content, filetype="pdf")
    # 每页输出项 (text, size, kind)：size 仅对 text 行有效；table/ocr 行 size=0，
    # 不参与字号统计与标题判定，整块作为正文
    raw_pages: list[list[tuple[str, float, str]]] = []
    with doc:
        # 1. 逐页提取：表格转 markdown、普通行排除表格区域、图片/扫描页 OCR
        for page in doc:
            table_bboxes: list[tuple] = []
            table_items: list[tuple[float, list[str]]] = []  # (y0, markdown 行列表)
            try:
                found = page.find_tables()
                for table in found.tables:
                    try:
                        md_lines = [line for line in table.to_markdown().split("\n") if line.strip()]
                    except Exception:
                        md_lines = []
                    if md_lines:
                        table_bboxes.append(table.bbox)
                        table_items.append((table.bbox[1], md_lines))
            except Exception as exc:
                logger.debug("pdf find_tables failed (ignored): %s", exc)

            data = page.get_text("dict", sort=True)
            line_items: list[tuple[float, list[tuple[str, float]]]] = []
            for block in data.get("blocks") or []:
                if block.get("type") != 0:  # 非文本块（图片等）
                    continue
                for line in block.get("lines") or []:
                    spans = [s for s in line.get("spans") or [] if s.get("text", "").strip()]
                    if not spans:
                        continue
                    if line.get("bbox") and _line_in_table(line["bbox"], table_bboxes):
                        continue  # 表格内部行，由 markdown 块替代
                    text = normalize_text("".join(s["text"] for s in spans))
                    if text:
                        size = round(max(s["size"] for s in spans), 1)
                        line_items.append((line["bbox"][1], [(text, size)]))

            # OCR：整页无文本层（扫描件）→ 整页渲染识别；有文本层 → 嵌入图片单独识别
            ocr_items: list[tuple[float, list[tuple[str, float]]]] = []
            if ocr is not None:
                if not line_items and not table_items:
                    texts = _ocr_page(page, ocr)
                    if texts:
                        logger.debug("pdf page ocr recovered %s lines", len(texts))
                        ocr_items.append((0.0, [(t, 0.0) for t in texts]))
                else:
                    for y0, texts in _ocr_embedded_images(doc, page, ocr):
                        ocr_items.append((y0, [(t, 0.0) for t in texts]))

            # 表格块 / OCR 块 / 普通行按页面位置排序，合并为页序列
            merged: list[tuple[float, list[tuple[str, float]], str]] = []
            merged += [(y0, [(text, 0.0) for text in md], _KIND_TABLE) for y0, md in table_items]
            merged += [(y0, lines, _KIND_OCR) for y0, lines in ocr_items]
            merged += [(y0, lines, _KIND_TEXT) for y0, lines in line_items]
            merged.sort(key=lambda item: item[0])
            page_items = [(text, size, kind) for _, lines, kind in merged for text, size in lines]
            if page_items:
                raw_pages.append(page_items)
    if not raw_pages:
        return []

    # 2. 行级页眉页脚/页码剔除（仅页面边缘的普通行；表格/OCR 行不参与）
    repeated = find_repeated_edge_lines(
        [[text for text, _, kind in page if kind == _KIND_TEXT] for page in raw_pages]
    )
    clean_pages: list[list[tuple[str, float, str]]] = []
    for page in raw_pages:
        kept: list[tuple[str, float, str]] = []
        n = len(page)
        for i, (text, size, kind) in enumerate(page):
            at_edge = i < _EDGE_LINES or i >= n - _EDGE_LINES
            if kind == _KIND_TEXT and at_edge and (text in repeated or _is_noise_line(text)):
                continue
            kept.append((text, size, kind))
        if kept:
            clean_pages.append(kept)

    # 3. 正文字号 = 普通行字号众数（平票取较小者：标题恒大于正文，避免小样本时
    #    众数落在大字号上导致标题漏判）；标题候选 = 显著大于正文的短行
    size_counts = Counter(
        size for page in clean_pages for _, size, kind in page if kind == _KIND_TEXT
    )
    if not size_counts:  # 全表格/全 OCR 页（无普通文本行），所有行整块作为正文
        agg = SectionAggregator()
        for page in clean_pages:
            for text, _, _ in page:
                agg.add_body(text)
        agg.flush()
        return agg.documents
    max_count = max(size_counts.values())
    body_size = min(size for size, cnt in size_counts.items() if cnt == max_count)
    heading_sizes = sorted(
        {
            size
            for page in clean_pages
            for _, size, kind in page
            if kind == _KIND_TEXT and size >= body_size * _HEADING_SIZE_RATIO
        },
        reverse=True,
    )
    # 字号降序聚类（差 ≤ 容差同档），前 3 档映射 h1/h2/h3
    clusters: list[float] = []
    for size in heading_sizes:
        if not clusters or clusters[-1] - size > _LEVEL_TOLERANCE:
            clusters.append(size)

    def _level_for(size: float) -> int:
        for idx, cluster in enumerate(clusters, start=1):
            if size >= cluster - _LEVEL_TOLERANCE:
                return min(idx, 3)
        return min(len(clusters), 3) or 1

    def _is_heading(text: str, size: float) -> bool:
        return (
            size >= body_size * _HEADING_SIZE_RATIO
            and len(text) <= _HEADING_MAX_LEN
            and not text.endswith(_TRAILING_PUNCT)
        )

    # 4. 分节聚合（与 docx 同一条路）；表格/OCR 行整块走正文
    agg = SectionAggregator()
    for page in clean_pages:
        for text, size, kind in page:
            if kind == _KIND_TEXT and _is_heading(text, size):
                agg.add_heading(text, _level_for(size))
            else:
                agg.add_body(text)
    agg.flush()
    return agg.documents


def _parse_pdf_plaintext(content: bytes) -> list[Document]:
    """回退：pypdf 逐页纯文本（无结构信息，切分阶段递归切分兜底）。"""
    try:
        from langchain_community.document_loaders import PyPDFLoader
    except Exception as exc:  # pragma: no cover - langchain-community 缺失时回退
        logger.warning("PyPDFLoader unavailable (%s), fallback to pypdf", exc)
        PyPDFLoader = None  # noqa: N806

    if PyPDFLoader is not None:
        fd, path = tempfile.mkstemp(suffix=".pdf")
        try:
            with os.fdopen(fd, "wb") as f:
                f.write(content)
            loaded = PyPDFLoader(path).load()
        finally:
            try:
                os.remove(path)
            except OSError:
                pass
        docs: list[Document] = []
        for idx, d in enumerate(loaded):
            meta = dict(d.metadata or {})
            page = meta.pop("page", None)
            page_number = page + 1 if isinstance(page, int) else idx + 1
            meta.pop("source", None)
            meta["page_number"] = page_number
            docs.append(Document(page_content=d.page_content or "", metadata=meta))
        return docs

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    return [
        Document(page_content=page.extract_text() or "", metadata={"page_number": i + 1})
        for i, page in enumerate(reader.pages)
    ]


# ----- docx -----

def _parse_docx(content: bytes) -> list[Document]:
    """docx：按标题分节聚合（SectionAggregator）——同一章节的段落与表格行
    （"表头: 值"）合并为一个 Document，标题文本作为章节首行进内容；
    长章节由切分阶段递归切分，子 chunk 经 split_documents 继承 title_path。
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from app.document.sections import SectionAggregator

    docx = DocxDocument(io.BytesIO(content))
    agg = SectionAggregator()

    def iter_blocks():
        """按 body 顺序迭代段落与表格（官方推荐的 element 遍历方式）。"""
        for child in docx.element.body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, docx)
            elif child.tag == qn("w:tbl"):
                yield Table(child, docx)

    for block in iter_blocks():
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            match = _HEADING_STYLE_RE.match(block.style.name) if block.style is not None else None
            if match:
                agg.add_heading(text, int(match.group(1)))
            else:
                agg.add_body(text)
        else:  # Table：行文本并入当前章节，保持表格上下文完整
            rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            if not rows:
                continue
            headers = rows[0]
            for row in rows[1:]:
                if not any(row):
                    continue
                text = "; ".join(f"{h}: {v}" for h, v in zip(headers, row) if v)
                if text:
                    agg.add_body(text)
    agg.flush()
    return agg.documents


# ----- xls/xlsx -----

def _parse_xlsx(content: bytes) -> list[Document]:
    """xlsx：逐工作表，首行表头，每行转 "表头1: 值1; 表头2: 值2..."，sheet 名入 metadata。

    注：openpyxl 不支持老格式 .xls（上传后解析会抛出并置 FAILED，提示转存 .xlsx）。
    """
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    docs: list[Document] = []
    try:
        for sheet in workbook.worksheets:
            headers: list[str] | None = None
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v).strip() for v in row]
                if not any(cells):
                    continue
                if headers is None:
                    headers = cells
                    continue
                text = "; ".join(f"{h}: {v}" for h, v in zip(headers, cells) if v)
                if text:
                    docs.append(
                        Document(page_content=text, metadata={"sheet": sheet.title})
                    )
    finally:
        workbook.close()
    return docs


# ----- txt/md -----

def _parse_txt(content: bytes) -> list[Document]:
    """txt：整体读入为单个 Document（切分阶段递归切分）。"""
    return [Document(page_content=_decode_text(content))]


def _parse_md(content: bytes) -> list[Document]:
    """md：保留原文单个 Document（切分阶段先用 MarkdownHeaderTextSplitter）。"""
    return [Document(page_content=_decode_text(content))]


_PARSERS = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "xlsx": _parse_xlsx,
    "txt": _parse_txt,
    "md": _parse_md,
}


def parse_file(content: bytes, filename: str) -> list[Document]:
    """解析器工厂：按扩展名解析文件字节为 langchain Document 列表。

    - .doc：抛 UnsupportedFormatError("不支持 .doc，请转存为 .docx")；
    - .xls：openpyxl 无法读取老格式，抛出异常由任务层置 FAILED；
    - 其他不支持格式：抛 UnsupportedFormatError。
    """
    ext = file_extension(filename)
    if ext == "doc":
        raise UnsupportedFormatError("不支持 .doc，请转存为 .docx")
    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedFormatError(f"不支持的文件格式：{ext or '未知'}")
    docs = parser(content)
    docs = clean_documents(docs, ext)  # 解析后统一清洗（规范化 + PDF 页眉页脚剔除）
    logger.info("parsed file=%s ext=%s blocks=%s", filename, ext, len(docs))
    return docs
